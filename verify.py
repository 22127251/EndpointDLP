"""Verify the generated corpus against the criteria by READING the files back.

For every output file this:
  * measures the on-disk size and checks it falls inside the file's size bracket;
  * extracts the text and counts VISA / CCCD / phone PII with anchored regexes
    (independent of the generator) and checks each type matches the expected count.

Run:  D:\\Code\\TestCaseGen\\.venv\\Scripts\\python.exe verify.py
"""
import re
import sys
import zipfile

from src import config

# --- PII counting patterns ---------------------------------------------------
# Digit lookarounds (not \b) so a value stays countable even if extraction glues
# it to an adjacent letter. The three shapes are mutually exclusive, so per-type
# counts never overlap (VISA starts with 4; CCCD is a bare 12-run; phone starts 0/+84).
RE_VISA = re.compile(r"(?<!\d)4\d{3} ?\d{4} ?\d{4} ?\d{4}(?!\d)")
RE_CCCD = re.compile(r"(?<!\d)\d{12}(?!\d)")
RE_PHONE = [
    re.compile(r"(?<!\d)0\d{2} \d{3} \d{4}(?!\d)"),   # grouped 096 123 4567
    re.compile(r"(?<!\d)\+84\d{9}(?!\d)"),            # intl  +84961234567
    re.compile(r"(?<!\d)0\d{9}(?!\d)"),               # plain 0961234567
]
WS = re.compile(r"\s+")


def count_pii(text: str) -> dict:
    text = WS.sub(" ", text)                          # heal wrapped/odd whitespace
    return {
        "visa": len(RE_VISA.findall(text)),
        "cccd": len(RE_CCCD.findall(text)),
        "phone": sum(len(p.findall(text)) for p in RE_PHONE),
    }


# --- Text extraction per format ---------------------------------------------
def _from_zip_xml(path, entry):
    with zipfile.ZipFile(path) as z:
        xml = z.read(entry).decode("utf-8", "ignore")
    return re.sub(r"<[^>]+>", " ", xml)               # tags -> spaces (keeps boundaries)


def extract_text(path, fmt: str) -> str:
    if fmt == "docx":
        from docx import Document
        d = Document(str(path))
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                parts.extend(c.text for c in row.cells)
        return " ".join(parts)
    if fmt == "xlsx":
        from openpyxl import load_workbook
        wb = load_workbook(str(path), read_only=True)
        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                parts.extend(str(v) for v in row if v is not None)
        wb.close()
        return " ".join(parts)
    if fmt in ("odt", "ods"):
        return _from_zip_xml(path, "content.xml")
    if fmt == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        return " ".join((pg.extract_text() or "") for pg in reader.pages)
    if fmt == "csv":
        return path.read_text(encoding="utf-8-sig", errors="ignore")
    return path.read_text(encoding="utf-8", errors="ignore")  # txt, md


# --- Main --------------------------------------------------------------------
FORMATS = ["docx", "odt", "pdf", "xlsx", "ods", "csv", "txt", "md"]


def main():
    print(f"{'file':16s} {'size (bytes)':>13s} {'size?':5s}  "
          f"{'VISA':>9s} {'CCCD':>9s} {'PHONE':>9s}  result")
    print("-" * 86)
    all_ok = True
    for bkey in config.BRACKET_ORDER:
        b = config.BRACKETS[bkey]
        exp = b["pii_per_type"]
        for fmt in FORMATS:
            path = config.OUTPUT_DIR / fmt / f"{fmt}_{bkey}.{fmt}"
            if not path.exists():
                print(f"{path.name:16s} {'MISSING':>13s}")
                all_ok = False
                continue
            size = path.stat().st_size
            size_ok = b["bracket_min"] < size < b["bracket_max"]
            c = count_pii(extract_text(path, fmt))
            pii_ok = c["visa"] == exp and c["cccd"] == exp and c["phone"] == exp
            ok = size_ok and pii_ok
            all_ok &= ok
            print(f"{path.name:16s} {size:>13,} {'OK' if size_ok else 'BAD':5s}  "
                  f"{c['visa']:>4d}/{exp:<4d} {c['cccd']:>4d}/{exp:<4d} {c['phone']:>4d}/{exp:<4d}  "
                  f"{'PASS' if ok else 'FAIL'}")
        print("-" * 86)
    print("ALL PASS" if all_ok else "SOME FAILED")
    return 0 if all_ok else 1


if __name__ == "__main__":
    sys.exit(main())
